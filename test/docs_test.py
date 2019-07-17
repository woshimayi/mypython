from docx import Document
from docx.shared import Inches





# 项目名称
# 项目版本号
# 当前时间
# 分支
# 编译命令
# Sf
# 问题列表
# 文件名
# MD5值
# sha256
# 服务器地址
# 编译人员
# 编译时间















# document = Document()

# document.add_heading('Document Title', 0)

# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')

# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )

# # document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# table = document.add_table(rows=1, cols=3, style='Medium Grid 1 Accent 1')

# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'

# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc

# document.add_page_break()

# document.save('demo.doc')


# 16 个信息


path = "demo.docx" #文件路径
document = Document(path) #读入文件
tables = document.tables #获取文件中的表格集

table0 = tables[0]#获取文件中的第一个表格
table1 = tables[1]
table2 = tables[2]
table3 = tables[3]
table4 = tables[4]
table5 = tables[5]
table6 = tables[6]
table7 = tables[7]
table8 = tables[8]
table9 = tables[9]
table10 = tables[10]
table11 = tables[11]

table0.cell(1,0).text = 'adfasdf'
print(table0.cell(1,0).text)

table10.cell(3,1).text = "123456789"
print(table10.cell(3,1).text)


# for i in range(1,len(table.rows)):#从表格第二行开始循环读取表格数据
#     result = table.cell(i,0).text+ " " +table.cell(i,1).text+table.cell(i,2).text
#     print(result)



document.save('demo.docx')